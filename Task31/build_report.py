from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from evidence_data import EVIDENCE, REFERENCES, ANNOTATED_IDS


ROOT = Path(__file__).resolve().parent
OUTDIR = ROOT / "outputs" / "task31_welfare_review_20260804"
OUTDIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUTDIR / "Asia_Pacific_Welfare_Losses_Review_2026.docx"
MD_PATH = ROOT / "review_manuscript.md"
FIGDIR = ROOT / "figures"

NAVY = "17365D"
TEAL = "087E8B"
GOLD = "D8A31A"
RED = "B5484D"
GREEN = "4B8F6A"
INK = "263238"
MUTED = "66727A"
PALE = "EFF4F6"
WHITE = "FFFFFF"
GRID = "CED9DF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=75, bottom=60, end=75):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    tr_pr.append(el)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_rows(row, val=True):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true" if val else "false")


def set_update_fields(doc):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_field(run, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_hyperlink(paragraph, text, url, color=TEAL):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_repeat_header(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(OxmlElement("w:tcFitText"))


def format_run(run, size=None, bold=None, italic=None, color=None, font="Aptos"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_rich_text(paragraph, text, size=10.3, color=INK):
    # Handles bold markers and URL tokens while preserving readable author-year text.
    token_re = re.compile(r"(\*\*.*?\*\*|https?://\S+)")
    pos = 0
    for m in token_re.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            format_run(r, size=size, color=color)
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            format_run(r, size=size, bold=True, color=color)
        else:
            clean = token.rstrip(".,;")
            suffix = token[len(clean):]
            add_hyperlink(paragraph, clean, clean)
            if suffix:
                r = paragraph.add_run(suffix)
                format_run(r, size=size, color=color)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        format_run(r, size=size, color=color)


def set_repeat_header_footer(section):
    section.header.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.clear()
    r = p.add_run("WELFARE LOSSES FROM AGGREGATE SHOCKS  |  ASIA AND THE PACIFIC")
    format_run(r, size=7.5, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(0)
    # bottom border on header paragraph
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), TEAL)
    pbdr.append(bottom)
    p_pr.append(pbdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    r = p.add_run("Working draft • Evidence cutoff 31 July 2026  |  ")
    format_run(r, size=7.2, color=MUTED)
    r = p.add_run()
    add_field(r, "PAGE")
    format_run(r, size=7.2, bold=True, color=NAVY)
    r = p.add_run(" / ")
    format_run(r, size=7.2, color=MUTED)
    r = p.add_run()
    add_field(r, "NUMPAGES")
    format_run(r, size=7.2, bold=True, color=NAVY)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(6.5)

    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 10),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 17, NAVY, 16, 7),
        ("Heading 2", 13, TEAL, 12, 5),
        ("Heading 3", 10.8, NAVY, 9, 3),
    ]:
        st = doc.styles[name]
        st.font.name = "Aptos Display" if name in ("Title", "Heading 1", "Heading 2") else "Aptos"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), st.font.name)
        st.font.size = Pt(size)
        st.font.bold = name != "Subtitle"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    doc.styles["Heading 1"].paragraph_format.page_break_before = True
    doc.styles["Heading 2"].paragraph_format.keep_with_next = True

    cap = doc.styles["Caption"]
    cap.font.name = "Aptos"
    cap.font.size = Pt(8.2)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    set_repeat_header_footer(section)
    set_update_fields(doc)


def add_cover(doc):
    # Hide first-page header/footer by creating a distinct cover section.
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("EVIDENCE REVIEW  •  ASIA AND THE PACIFIC")
    format_run(r, size=9, bold=True, color=TEAL)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("Welfare Losses from\nAggregate Socioeconomic and\nEnvironmental Shocks")
    format_run(r, size=28, bold=True, color=NAVY, font="Aptos Display")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("A multidimensional review of COVID-19, economic crises, disasters, environmental degradation, and climate change")
    format_run(r, size=13.2, color=MUTED)

    # Cover rule
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [3.2, 1.6, 1.6]
    fills = [NAVY, TEAL, GOLD]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Inches(widths[i])
        set_cell_shading(cell, fills[i])
        cell.text = ""
    table.rows[0].height = Inches(0.13)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Prepared as a near-publication-ready chapter for internal review")
    format_run(r, size=10.5, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Economic Research and Development Impact Department")
    format_run(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Evidence cutoff: 31 July 2026")
    format_run(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    r = p.add_run("Review corpus: 52 quantitative studies and assessments")
    format_run(r, size=10, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    r = p.add_run("Scope note")
    format_run(r, size=9.3, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.55)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Observed losses, causal estimates, accounting assessments, exposure measures, and long-run scenarios are labelled separately. Estimates are not summed across incompatible units or overlapping populations.")
    format_run(r, size=9, color=MUTED)
    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run()
    add_field(r, 'TOC \\o "1-2" \\h \\z \\u')
    p = doc.add_paragraph()
    r = p.add_run("The table of contents updates when the document opens in Microsoft Word.")
    format_run(r, size=8.2, italic=True, color=MUTED)
    doc.add_page_break()


def table_font(table, size=7.0):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    format_run(r, size=size, color=INK)


def add_table(doc, headers, rows, widths=None, font_size=6.8, caption=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[i])
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                format_run(r, size=font_size, bold=True, color=WHITE)
    for ridx, values in enumerate(rows):
        row = table.add_row()
        cant_split(row)
        for i, value in enumerate(values):
            cell = row.cells[i]
            if widths:
                cell.width = Inches(widths[i])
            if ridx % 2 == 1:
                set_cell_shading(cell, PALE)
            cell.text = str(value)
    table_font(table, font_size)
    # Re-apply header styling after table-wide font.
    for cell in table.rows[0].cells:
        set_cell_shading(cell, NAVY)
        for p in cell.paragraphs:
            for r in p.runs:
                format_run(r, size=font_size, bold=True, color=WHITE)
    if caption:
        p = doc.add_paragraph(style="Caption")
        p.add_run(caption)
    return table


def rows_for(ids):
    by_id = {e["id"]: e for e in EVIDENCE}
    return [by_id[i] for i in ids]


def add_required_table(doc, number):
    if number == "1":
        ids = ["C01", "C02", "C03", "C04", "C05", "C06", "C07", "C08", "C09", "C10", "C11", "C12", "C14", "C15"]
        es = rows_for(ids)
        rows = [[e["study"], e["geography"], e["population"], e["welfare_indicator"], e["estimate"], e["methodology"]] for e in es]
        add_table(doc, ["Study", "Geography", "Population", "Welfare indicator", "Estimated loss", "Methodology"], rows,
                  widths=[0.72, 1.05, 0.95, 1.05, 1.75, 1.25], font_size=6.15,
                  caption="Table 1. Major studies on COVID-19 welfare losses in Asia and the Pacific. Source: authors' synthesis of cited studies.")
    elif number == "2":
        es = [e for e in EVIDENCE if e["category"] == "Economic shock"]
        rows = [[e["study"], e["shock"], e["geography"], e["welfare_indicator"], e["estimate"]] for e in es]
        add_table(doc, ["Study", "Shock type", "Geography", "Welfare measure", "Magnitude"], rows,
                  widths=[0.92, 1.18, 1.18, 1.22, 2.20], font_size=6.35,
                  caption="Table 2. Major studies on economic shocks, 2015-present. Source: authors' synthesis of cited studies.")
    elif number == "3":
        ids = ["N01", "N02", "N03", "N04", "N06", "N07", "N08", "N09", "N10", "N12", "N13", "N14", "N15", "N16", "N17", "N18", "N19", "N21", "N22", "N23", "N24"]
        es = rows_for(ids)
        rows = [[e["study"], e["shock"], e["geography"], e["welfare_indicator"], e["estimate"]] for e in es]
        add_table(doc, ["Study", "Hazard type", "Geography", "Welfare outcome", "Magnitude of loss"], rows,
                  widths=[0.88, 1.05, 1.18, 1.28, 2.31], font_size=6.15,
                  caption="Table 3. Major studies on environmental and climate shocks. Exposure and scenario estimates are labelled and are not treated as realized loss.")
    elif number == "4":
        rows = [
            ["Children (0-17)", "Learning, nutrition, service disruption, protection and psychosocial development", "Learning poverty 60%→78% in South Asia; future earnings −up to 14.4%; wasting risk +9% after a 5% real food-price increase", "South Asia; disaster-prone Southeast Asia and Pacific; pastoral Central/East Asia"],
            ["Working-age adults (18-64)", "Jobs, hours, earnings, informality, migration, debt, care and mental health", "81 million Asia-Pacific jobs lost in 2020; women +8 pp work stoppage; Lao real wages about −33% in 2023", "Region-wide; tourism-dependent Southeast Asia and Pacific; crisis economies"],
            ["Older persons (65+)", "Mortality, chronic disease, care access, isolation, fixed-income erosion and evacuation", "Pre-vaccine IFR 8.29% at 80+; pandemic mortality concentrated at older ages; non-fatal Asia-specific estimates sparse", "All subregions; highest gaps in low-capacity and conflict settings"],
        ]
        add_table(doc, ["Group", "Primary welfare impacts", "Quantitative estimates", "Regions most affected"], rows,
                  widths=[1.05, 1.65, 2.45, 1.75], font_size=6.8,
                  caption="Table 4. Welfare losses by demographic group.")
    elif number == "5":
        rows = [
            ["COVID-19", "Very high and region-wide: GDP, jobs, poverty, mortality, food security and services", "High: excess mortality, learning, nutrition, debt and scarring", "Large regional models, surveys, mortality and learning evidence", "High for direction and ranking; medium for exact totals"],
            ["Inflation and food/energy prices", "Moderate regionally; very high for poor net consumers and crisis countries", "Medium-high when nutrition, schooling or assets deteriorate", "Good microsimulation; strong child-nutrition micro evidence", "High for regressivity/mechanism; medium for national totals"],
            ["Macro, debt, currency and trade shocks", "Highly heterogeneous; severe in Sri Lanka, Afghanistan, Myanmar and Lao PDR", "High if investment, services, female work and institutions remain impaired", "Good crisis monitoring; attribution usually weak", "Medium; low in conflict/data-collapse settings"],
            ["Disasters and environmental degradation", "Extreme locally; event effects up to 31%-64% of GDP in Pacific cases", "High through assets, displacement, health, nutrition and repeated exposure", "Strong PDNAs and causal health/productivity evidence; welfare aggregation sparse", "High for event damage and key mechanisms"],
            ["Long-run climate change", "Already visible in heat mortality, productivity and disaster risk", "Potentially the largest cumulative burden; ADB high-end GDP gap 41% by 2100", "Multiple models agree on direction; magnitude varies with adaptation and damage functions", "Medium for broad order; low for distant point estimates"],
        ]
        add_table(doc, ["Shock category", "Short-term welfare losses", "Long-term welfare losses", "Evidence strength", "Confidence assessment"], rows,
                  widths=[1.00, 1.50, 1.55, 1.45, 1.45], font_size=6.25,
                  caption="Table 5. Comparative ranking of shocks. Rankings are ordinal and preserve differences in unit, horizon, and counterfactual.")
    elif number == "KEY":
        ids = ["C01", "C02", "C03", "C12", "C08", "E01", "E05", "E06", "E10", "N14", "N16", "N17", "N10", "N21", "N02"]
        es = rows_for(ids)
        rows = [[e["study"], e["estimate"], e["evidence_type"], e["confidence"]] for e in es]
        add_table(doc, ["Study", "Key estimate", "Estimate type", "Confidence"], rows,
                  widths=[1.05, 3.65, 1.35, 0.75], font_size=6.55,
                  caption="Key quantitative estimates. Units, baselines and horizons follow the original sources; estimates are not additive.")


FIGS = {
    "1": ("figure_1_conceptual_pathways.png", "Figure 1. Conceptual pathways linking aggregate shocks to welfare losses."),
    "2": ("figure_2_comparative_magnitude.png", "Figure 2. Comparative magnitude of selected welfare losses. Different panels retain different units."),
    "3": ("figure_3_lifecycle_impacts.png", "Figure 3. Life-cycle impacts among children, working-age adults, and older persons."),
    "4": ("figure_4_geographic_distribution.png", "Figure 4. Geographic distribution of reviewed evidence and representative losses across Asia-Pacific subregions."),
}


def add_figure(doc, number):
    filename, caption = FIGS[number]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(FIGDIR / filename), width=Inches(6.75))
    p = doc.add_paragraph(style="Caption")
    p.add_run(caption)
    p = doc.add_paragraph(style="Caption")
    p.add_run("Source: authors' synthesis of the study-level evidence register.")


def add_method_callout(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
    p = cell.paragraphs[0]
    r = p.add_run("Reading the estimates")
    format_run(r, size=9.2, bold=True, color=NAVY)
    p = cell.add_paragraph()
    add_rich_text(p, "Ranges preserve the original scenario or uncertainty interval. GDP gaps, asset losses, poverty changes, deaths, exposure counts, and lifetime earnings are not summed. Confidence applies to the estimate's use in this review.", size=8.6, color=INK)


def add_main_text(doc):
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    started = False
    abstract_started = False
    first_h1 = True
    for raw in lines:
        line = raw.strip()
        if not started:
            if line == "## Abstract":
                started = True
                abstract_started = True
                p = doc.add_paragraph("Abstract", style="Heading 1")
                p.paragraph_format.page_break_before = False
            continue
        if not line:
            continue
        if line.startswith("<FIGURE:"):
            add_figure(doc, line.split(":")[1].rstrip(">"))
            continue
        if line.startswith("<TABLE:"):
            add_required_table(doc, line.split(":")[1].rstrip(">"))
            continue
        if line.startswith("## "):
            text = line[3:]
            p = doc.add_paragraph(text, style="Heading 1")
            if abstract_started:
                abstract_started = False
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:], style="Heading 2")
            continue
        if line.startswith("#### "):
            p = doc.add_paragraph(line[5:], style="Heading 3")
            continue
        if line.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            add_rich_text(p, line, size=8.8, color=MUTED)
            add_method_callout(doc)
            continue
        p = doc.add_paragraph()
        if line.startswith("**Figure"):
            p.style = doc.styles["Caption"]
            add_rich_text(p, line, size=8.3, color=MUTED)
        else:
            add_rich_text(p, line, size=10.3, color=INK)


def add_references(doc):
    p = doc.add_paragraph("References", style="Heading 1")
    refs = sorted(dict.fromkeys(REFERENCES), key=lambda x: re.sub(r"^[^A-Za-z]+", "", x).lower())
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(4.2)
        p.paragraph_format.line_spacing = 1.05
        add_rich_text(p, ref, size=8.3, color=INK)


def add_annotated_bibliography(doc):
    p = doc.add_paragraph("Annotated bibliography: 20 influential studies", style="Heading 1")
    intro = doc.add_paragraph()
    add_rich_text(intro, "The studies below were selected for regional reach, empirical credibility, methodological influence, or importance to cross-shock comparison. Each note reports the study's quantitative contribution, design, interpretive value, and principal limitation.", size=9.5)
    by_id = {e["id"]: e for e in EVIDENCE}
    for n, sid in enumerate(ANNOTATED_IDS, start=1):
        e = by_id[sid]
        p = doc.add_paragraph(style="Heading 3")
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"{n}. {e['study']} — {e['source']}")
        format_run(r, size=10.2, bold=True, color=NAVY)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_rich_text(p, f"Coverage and finding. {e['geography']}; {e['population']}. The study reports {e['estimate']}.", size=8.9)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_rich_text(p, f"Method and identification. {e['methodology']}. The identifying basis is {e['identification'].lower()}.", size=8.9)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_rich_text(p, f"Use and caution. It is influential here because it supplies a {e['evidence_type'].lower()} for {e['welfare_indicator'].lower()}. Principal limitation: {e['limitations']}. Confidence: {e['confidence']}.", size=8.9)
        p.add_run("  ")
        add_hyperlink(p, "Source", e["url"])


def add_reproducibility_note(doc):
    p = doc.add_paragraph("Review protocol and reproducibility note", style="Heading 1")
    for text in [
        "The companion workbook contains one row per quantitative source, the required extraction fields, confidence coding, source URLs, and the data used for Figures 2-4. The manuscript and figures were generated from the same structured register so that numeric updates can be traced.",
        "This is a structured rapid evidence review with systematic-scoping elements, not a registered systematic review. Before external journal submission, the search should be refreshed in Scopus, Web of Science, EconLit, PubMed, and the ADB Library; titles, abstracts, and full texts should be dual-screened; and a second reviewer should independently verify all extracted estimates.",
        "Version: 2026-08-04. Evidence cutoff: 2026-07-31. Main narrative word count: approximately 9,655 words, excluding references, tables, and the annotated bibliography."
    ]:
        p = doc.add_paragraph()
        add_rich_text(p, text, size=9.5)


def add_doc_properties(doc):
    props = doc.core_properties
    props.title = "Welfare Losses from Aggregate Socioeconomic and Environmental Shocks in Asia and the Pacific"
    props.subject = "Publication-ready multidimensional literature review"
    props.author = "Economic Research and Development Impact Department"
    props.keywords = "Asia Pacific, welfare loss, COVID-19, inflation, disasters, climate change, poverty, human capital"
    props.comments = "Working draft for internal review; evidence cutoff 31 July 2026."


def build():
    doc = Document()
    style_document(doc)
    add_doc_properties(doc)
    add_cover(doc)
    add_toc(doc)
    add_main_text(doc)
    add_references(doc)
    add_annotated_bibliography(doc)
    add_reproducibility_note(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()

